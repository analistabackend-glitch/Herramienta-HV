"""
extractor.py
============
Extrae texto de PDFs y DOCXs con manejo mejorado de:
- PDFs con layout de 2 columnas (muy común en HVs colombianas)
- Sistema de caché para evitar re-extracción
- Fallback a extracción simple si el modo columnas falla

Instalar: pip install pdfplumber python-docx
"""

from docx import Document
from pathlib import Path
import hashlib
import os

# Carpeta de caché para textos extraídos
CACHE_DIR = Path("cache_textos")
CACHE_DIR.mkdir(exist_ok=True)


def _cache_path(file_path):
    """Genera ruta de caché única basada en el hash del archivo."""
    file_path = Path(file_path)
    try:
        with open(file_path, 'rb') as f:
            file_hash = hashlib.md5(f.read(8192)).hexdigest()
    except:
        file_hash = str(int(file_path.stat().st_mtime))
    return CACHE_DIR / f"{file_path.stem}_{file_hash}.txt"


def _extraer_pagina_con_columnas(page) -> str:
    """
    Extrae texto de una página intentando respetar el orden visual.
    
    Estrategia:
    1. Intenta extracción normal primero
    2. Si detecta layout de columnas (palabras muy separadas horizontalmente),
       ordena los bloques de texto por posición (top→bottom, left→right)
    """
    # Intento 1: extracción estándar
    texto_simple = page.extract_text(x_tolerance=3, y_tolerance=3)
    
    # Intento 2: extracción por palabras ordenadas por posición
    # Esto resuelve el problema de columnas mezcladas
    try:
        words = page.extract_words(
            x_tolerance=5,
            y_tolerance=5,
            keep_blank_chars=False,
            use_text_flow=True  # ← clave: respeta el flujo visual del texto
        )
        
        if not words:
            return texto_simple or ""
        
        # Detectar si hay columnas: si el ancho de la página se usa en zonas muy separadas
        x_positions = [w["x0"] for w in words]
        page_width = page.width
        
        # Reconstruir texto agrupando palabras en líneas por proximidad vertical
        lineas = {}
        for word in words:
            # Agrupar por línea (tolerancia de 5px en Y)
            y_key = round(word["top"] / 5) * 5
            if y_key not in lineas:
                lineas[y_key] = []
            lineas[y_key].append(word)
        
        # Ordenar líneas de arriba a abajo, palabras de izquierda a derecha
        texto_reconstruido = []
        for y_key in sorted(lineas.keys()):
            palabras_linea = sorted(lineas[y_key], key=lambda w: w["x0"])
            texto_reconstruido.append(" ".join(w["text"] for w in palabras_linea))
        
        texto_por_palabras = "\n".join(texto_reconstruido)
        
        # Usar el más largo (más contenido extraído)
        if len(texto_por_palabras) > len(texto_simple or ""):
            return texto_por_palabras
        
        return texto_simple or ""
        
    except Exception:
        return texto_simple or ""


def leer_pdf(path):
    """
    Extrae texto de PDF con manejo de columnas múltiples.
    Muy común en HVs colombianas con layout de 2 columnas.
    """
    texto = ""

    try:
        import pdfplumber

        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                t = _extraer_pagina_con_columnas(page)
                if t:
                    texto += t + "\n"

        return texto

    except Exception as e:
        print(f"  ❌ Error extrayendo PDF {os.path.basename(path)}: {e}")
        return ""


def leer_docx(path):
    """Extrae texto de archivo DOCX."""
    try:
        doc = Document(path)
        partes = []
        
        # Párrafos normales
        for p in doc.paragraphs:
            if p.text.strip():
                partes.append(p.text)
        
        # Tablas (muchas HVs usan tablas para el layout)
        for tabla in doc.tables:
            for fila in tabla.rows:
                for celda in fila.cells:
                    if celda.text.strip():
                        partes.append(celda.text.strip())
        
        return "\n".join(partes)
    except Exception as e:
        print(f"  ❌ Error extrayendo DOCX {os.path.basename(path)}: {e}")
        return ""


def extraer_texto(path):
    """
    Extrae texto de PDF o DOCX con sistema de caché.
    """
    cache_file = _cache_path(path)

    # Intentar leer del caché
    if cache_file.exists():
        try:
            texto_cache = cache_file.read_text(encoding='utf-8')
            if texto_cache.strip():
                return texto_cache
        except Exception as e:
            print(f"  ⚠️  Error leyendo caché para {os.path.basename(path)}: {e}")

    # Extraer texto del archivo original
    if path.endswith(".pdf"):
        texto = leer_pdf(path)
    elif path.endswith(".docx"):
        texto = leer_docx(path)
    else:
        return ""

    # Guardar en caché
    if texto.strip():
        try:
            cache_file.write_text(texto, encoding='utf-8')
        except Exception as e:
            print(f"  ⚠️  No se pudo guardar caché para {os.path.basename(path)}: {e}")

    return texto


def limpiar_cache():
    """Limpia archivos de caché."""
    try:
        import shutil
        if CACHE_DIR.exists():
            shutil.rmtree(CACHE_DIR)
            CACHE_DIR.mkdir()
            print("✅ Caché de textos limpiado")
    except Exception as e:
        print(f"⚠️  Error limpiando caché: {e}")


if __name__ == "__main__":
    import time
    import sys

    if len(sys.argv) < 2:
        print("Uso: python extractor.py <archivo.pdf>")
        sys.exit(1)

    archivo = sys.argv[1]

    limpiar_cache()
    inicio = time.time()
    texto1 = extraer_texto(archivo)
    tiempo1 = time.time() - inicio

    inicio = time.time()
    texto2 = extraer_texto(archivo)
    tiempo2 = time.time() - inicio

    print(f"\n{'='*60}")
    print(f"Archivo: {os.path.basename(archivo)}")
    print(f"Tamaño: {len(texto1):,} caracteres")
    print(f"Primera extracción (sin caché): {tiempo1:.3f} seg")
    print(f"Segunda extracción (con caché): {tiempo2:.3f} seg")
    print(f"Aceleración: {tiempo1/tiempo2:.1f}x más rápido")
    print(f"\nPrimeros 500 caracteres del texto extraído:")
    print(f"{'='*60}")
    print(texto1[:500])
